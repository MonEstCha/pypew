import os
from datetime import datetime
from functools import lru_cache
from typing import Optional, List

import pandas as pd
from attr import define, field
from docx import Document

from utils import get_neh_df

feasts_fields = ['name', 'introit', 'collect', 'epistle_ref', 'epistle',
                 'gat', 'gradual', 'alleluia', 'tract', 'gospel_ref',
                 'gospel', 'offertory', 'communion']


class NotFoundError(Exception):
    pass


class MultipleReturnedError(Exception):
    pass


def get(collection, **kwargs):
    f = lambda x: all(getattr(x, k) == v for k, v in kwargs.items())
    filtered = list(filter(f, collection))

    n = len(filtered)
    if n == 0:
        raise NotFoundError(kwargs)
    if n > 1:
        raise MultipleReturnedError(kwargs, n)
    return filtered[0]


@define
class Feast:
    name: str = field()
    introit: str = field()
    collect: str = field()
    epistle_ref: str = field()
    epistle: str = field()
    gat: str = field()
    gradual: str = field()
    alleluia: str = field()
    tract: str = field()
    gospel_ref: str = field()
    gospel: str = field()
    offertory: str = field()
    communion: str = field()

    feasts_df = pd.read_csv(
        os.path.join(os.path.dirname(__file__), 'data', 'feasts.csv')
    )
    assert list(feasts_df.columns) == feasts_fields

    @classmethod
    def all(cls):
        # attrs provides a __init__ that takes kwargs
        # noinspection PyArgumentList
        return [cls(**info) for _, info in cls.feasts_df.iterrows()]

    @classmethod
    def get(cls, **kwargs):
        return get(cls.all(), **kwargs)

    def create_docx(self, path):
        document = Document()
        document.add_heading(self.name, 0)
        document.save(path)


@define
class Music:
    hymns_df = get_neh_df()

    title: str = field()
    category: str = field()  # Anthem or Hymn or Plainsong
    composer: Optional[str] = field()
    lyrics: Optional[str] = field()
    ref: Optional[str] = field()

    @classmethod
    def neh_hymns(cls) -> List['Music']:
        return [
            Music(
                title=record.firstLine,
                category='Hymn',
                composer=None,
                lyrics=None,
                ref=f'NEH: {record.number}'
            ) for record in get_neh_df().itertuples()
        ]

    @classmethod
    def get_neh_hymn_by_ref(cls, ref: str) -> 'Music':
        return next(filter(lambda h: h.ref == ref, cls.neh_hymns()))

    def __str__(self):
        if self.category == 'Hymn':
            return f'{self.ref}, {self.title}'
        return super().__str__()


@define
class Service:
    title: str = field()
    date: datetime.date = field()
    celebrant: str = field()
    preacher: str = field()
    primary_feast: Feast = field()
    secondary_feast: Optional[Feast] = field()
    introit_hymn: Optional[Music] = field()
    offertory_hymn: Optional[Music] = field()
    recessional_hymn: Optional[Music] = field()
    anthem: Music = field()

    @property
    def collects(self) -> List[str]:
        out = []
        if self.primary_feast.collect:
            out.append(self.primary_feast.collect)
        if self.secondary_feast and self.secondary_feast.collect:
            out.append(self.secondary_feast.collect)
        return out

    # TODO primary or secondary?
    @property
    def introit_proper(self) -> str:
        return self.primary_feast.introit

    @property
    def offertory_proper(self) -> str:
        return self.primary_feast.offertory

    @property
    def communion_proper(self) -> str:
        return self.primary_feast.communion

    @property
    def epistle_ref(self) -> str:
        return self.primary_feast.epistle_ref

    @property
    def epistle(self) -> str:
        return self.primary_feast.epistle

    @property
    def gospel_ref(self) -> str:
        return self.primary_feast.gospel_ref

    @property
    def gospel(self) -> str:
        return self.primary_feast.gospel

    def create_docx(self, path):
        import filters  # local import to avoid circular import

        document = Document()
        document.add_heading(self.title, 0)
        p = document.add_paragraph('')
        if self.date:
            run = p.add_run(filters.english_date(self.date))
            run.italic = True
        (filters.english_date(self.date))
        document.save(path)
